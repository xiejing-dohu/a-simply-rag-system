"""知识库运行时处理服务模块

包含文档提取（支持 .docx, .xlsx, .csv, .txt, .md, .pdf 等格式）、
基于 Tiktoken 的 Token 切片、Embedding 向量化提取 API 调用、
以及 Milvus Collection 的创建、查询、删除与文档切片 Upsert/Delete 操作。
"""

from __future__ import annotations

import io
import hashlib
import os
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

import httpx
import pandas as pd
import tiktoken
from tenacity import AsyncRetrying, retry_if_exception_type, stop_after_attempt, wait_exponential
from docx import Document as DocxDocument
from pypdf import PdfReader
from pymilvus import (
    Collection,
    DataType,
    Function,
    FunctionType,
    MilvusClient,
    connections,
    utility,
)

# 支持上传解析的文件后缀集合
SUPPORTED_EXTENSIONS = {".docx", ".xlsx", ".xls", ".csv", ".txt", ".md", ".pdf"}
# DashScope 推荐支持的向量维度
DASHSCOPE_DIMENSIONS = [64, 128, 256, 512, 768, 1024, 1536, 2048]


class KnowledgeProcessingError(RuntimeError):
    """知识库处理基础异常类"""
    pass


class EmbeddingServiceError(KnowledgeProcessingError):
    """Embedding 服务生成向量失败异常"""
    pass


class TransientEmbeddingError(RuntimeError):
    """Embedding 暂态网络/服务不可用异常"""
    pass


class LegacyCollectionSchemaError(KnowledgeProcessingError):
    """旧版无法保证幂等的 Milvus Schema 异常"""
    pass


def embedding_config() -> dict[str, Any]:
    """获取当前系统配置的 Embedding 模型与维度支持信息

    Returns:
        dict[str, Any]: 模型元数据字典
    """
    model = os.getenv("EMBEDDING_MODEL", "").strip()
    api_base = os.getenv("OPENAI_API_BASE", "").strip()
    dimensions = DASHSCOPE_DIMENSIONS if "dashscope" in api_base.lower() else [256, 512, 768, 1024, 1536, 3072]
    return {
        "model": model,
        "provider": "DashScope" if "dashscope" in api_base.lower() else "OpenAI Compatible",
        "supported_dimensions": dimensions,
        "default_dimension": 1024 if "dashscope" in api_base.lower() else 1536,
        "supported_extensions": sorted(SUPPORTED_EXTENSIONS),
    }


def _decode_text(content: bytes) -> str:
    """尝试以 UTF-8-SIG, UTF-8, GB18030 等多种编码方式解码文本字节流"""
    for encoding in ("utf-8-sig", "utf-8", "gb18030"):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise KnowledgeProcessingError("文本编码无法识别，请转换为 UTF-8 后重试")


def _stringify_cell(value: Any) -> str:
    """将 Excel 单元格转换为干净的文本串"""
    if pd.isna(value):
        return ""
    if isinstance(value, float) and value.is_integer():
        return str(int(value))
    return str(value).strip()


def _extract_excel(content: bytes) -> str:
    """提取 Excel (.xlsx, .xls) 中所有 Sheet 表格的结构化文本"""
    try:
        workbook = pd.ExcelFile(io.BytesIO(content))
    except Exception as exc:
        raise KnowledgeProcessingError(f"Excel 文件无法读取: {exc}") from exc

    sections: list[str] = []
    for sheet_name in workbook.sheet_names:
        frame = workbook.parse(sheet_name=sheet_name, dtype=object)
        frame = frame.dropna(how="all").dropna(axis=1, how="all")
        if frame.empty:
            continue
        columns = [str(column).strip() or f"列{index + 1}" for index, column in enumerate(frame.columns)]
        rows = [f"[工作表: {sheet_name}]", "字段: " + " | ".join(columns)]
        for row_number, values in enumerate(frame.itertuples(index=False, name=None), start=2):
            parts = [
                f"{column}: {_stringify_cell(value)}"
                for column, value in zip(columns, values)
                if _stringify_cell(value)
            ]
            if parts:
                rows.append(f"第{row_number}行 | " + " | ".join(parts))
        if len(rows) > 2:
            sections.append("\n".join(rows))
    if not sections:
        raise KnowledgeProcessingError("Excel 文件中没有可处理的数据")
    return "\n\n".join(sections)


def extract_document(file_name: str, content: bytes) -> tuple[str, str]:
    """主文档内容提取函数

    根据文件后缀（TXT, MD, CSV, DOCX, XLSX, PDF）调用不同解析器提取纯文本。

    Args:
        file_name (str): 文件名称
        content (bytes): 文件二进制流

    Returns:
        tuple[str, str]: (提取出的纯文本, 文件来源类型名)
    """
    suffix = Path(file_name).suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        if suffix == ".doc":
            raise KnowledgeProcessingError("暂不支持旧版 .doc，请另存为 .docx 后上传")
        raise KnowledgeProcessingError(f"不支持 {suffix or '无扩展名'} 文件")

    if suffix in {".txt", ".md", ".csv"}:
        text = _decode_text(content)
    elif suffix in {".xlsx", ".xls"}:
        text = _extract_excel(content)
    elif suffix == ".docx":
        try:
            document = DocxDocument(io.BytesIO(content))
            blocks = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
            for table_index, table in enumerate(document.tables, start=1):
                blocks.append(f"[表格 {table_index}]")
                blocks.extend(
                    " | ".join(cell.text.strip() for cell in row.cells)
                    for row in table.rows
                    if any(cell.text.strip() for cell in row.cells)
                )
            text = "\n".join(blocks)
        except Exception as exc:
            raise KnowledgeProcessingError(f"DOCX 文件无法读取: {exc}") from exc
    else:
        try:
            reader = PdfReader(io.BytesIO(content))
            text = "\n\n".join(page.extract_text() or "" for page in reader.pages)
        except Exception as exc:
            raise KnowledgeProcessingError(f"PDF 文件无法读取: {exc}") from exc

    text = text.replace("\x00", "").strip()
    if not text:
        raise KnowledgeProcessingError("文件中没有提取到可处理文本")
    return text, suffix.removeprefix(".")


def split_text_by_tokens(text: str, chunk_tokens: int, overlap_tokens: int) -> list[dict[str, Any]]:
    """使用 cl100k_base 编码器按指定的 Token 大小与 Overlap 滑动窗口切片文本

    Args:
        text (str): 待切片纯文本
        chunk_tokens (int): 单个 Chunk 目标 Token 数量
        overlap_tokens (int): 相邻 Chunk 间的重叠 Token 数量

    Returns:
        list[dict[str, Any]]: 切片字典列表 [{"text": "...", "token_count": 128}]
    """
    if not 64 <= chunk_tokens <= 8192:
        raise KnowledgeProcessingError("切片 Token 数必须在 64 到 8192 之间")
    if overlap_tokens < 0 or overlap_tokens >= chunk_tokens:
        raise KnowledgeProcessingError("Overlap 必须大于等于 0 且小于切片 Token 数")

    encoding = tiktoken.get_encoding("cl100k_base")
    token_ids = encoding.encode(text)
    if not token_ids:
        raise KnowledgeProcessingError("文件没有可切片的 Token")

    chunks: list[dict[str, Any]] = []
    step = chunk_tokens - overlap_tokens
    for start in range(0, len(token_ids), step):
        current_ids = token_ids[start : start + chunk_tokens]
        chunk_text = encoding.decode(current_ids).strip()
        if chunk_text:
            chunks.append({"text": chunk_text, "token_count": len(current_ids)})
        if start + chunk_tokens >= len(token_ids):
            break
    return chunks


def connect_milvus() -> None:
    """建立与 Milvus 向量数据库的连接"""
    connections.connect(
        alias="default",
        host=os.getenv("MILVUS_HOST", "localhost"),
        port=os.getenv("MILVUS_PORT", "19530"),
    )


def create_collection(collection_name: str, vector_dimension: int) -> None:
    """在 Milvus 中创建包含 Dense 向量与 BM25 稀疏向量的集合与索引结构

    Args:
        collection_name (str): 集合名称
        vector_dimension (int): 密集向量维度
    """
    connect_milvus()
    if utility.has_collection(collection_name):
        return
    uri = (
        f"http://{os.getenv('MILVUS_HOST', 'localhost')}:"
        f"{os.getenv('MILVUS_PORT', '19530')}"
    )
    client = MilvusClient(uri=uri)
    schema = MilvusClient.create_schema(auto_id=False, enable_dynamic_field=False)
    schema.add_field(
        field_name="id", datatype=DataType.INT64, is_primary=True, auto_id=False
    )
    schema.add_field(
        field_name="document_id", datatype=DataType.VARCHAR, max_length=36
    )
    schema.add_field(
        field_name="text",
        datatype=DataType.VARCHAR,
        max_length=65535,
        enable_analyzer=True,
        analyzer_params={"tokenizer": "standard", "filter": ["lowercase"]},
    )
    schema.add_field(field_name="sparse", datatype=DataType.SPARSE_FLOAT_VECTOR)
    schema.add_field(
        field_name="document_name", datatype=DataType.VARCHAR, max_length=1024
    )
    schema.add_field(
        field_name="source_type", datatype=DataType.VARCHAR, max_length=32
    )
    schema.add_field(field_name="chunk_index", datatype=DataType.INT64)
    schema.add_field(field_name="token_count", datatype=DataType.INT64)
    schema.add_field(field_name="uploaded_by", datatype=DataType.INT64)
    schema.add_field(
        field_name="created_at", datatype=DataType.VARCHAR, max_length=64
    )
    schema.add_field(
        field_name="embedding",
        datatype=DataType.FLOAT_VECTOR,
        dim=vector_dimension,
    )
    # 添加自动基于 text 生成 BM25 稀疏向量的 Function
    schema.add_function(
        Function(
            name="text_bm25",
            input_field_names=["text"],
            output_field_names=["sparse"],
            function_type=FunctionType.BM25,
        )
    )
    indexes = MilvusClient.prepare_index_params()
    indexes.add_index(
        field_name="embedding",
        metric_type="COSINE",
        index_type="IVF_FLAT",
        params={"nlist": 128},
    )
    indexes.add_index(
        field_name="sparse",
        metric_type="BM25",
        index_type="SPARSE_INVERTED_INDEX",
        params={"inverted_index_algo": "DAAT_MAXSCORE"},
    )
    client.create_collection(
        collection_name=collection_name,
        schema=schema,
        index_params=indexes,
    )


def drop_collection(collection_name: str) -> None:
    """删除指定的 Milvus 集合"""
    connect_milvus()
    if utility.has_collection(collection_name):
        utility.drop_collection(collection_name)


def collection_exists(collection_name: str) -> bool:
    """检查 Milvus 集合是否存在"""
    connect_milvus()
    return bool(utility.has_collection(collection_name))


def require_idempotent_collection(collection_name: str) -> None:
    """校验集合 Schema 是否支持幂等主键索引，拒绝自动 ID 的旧版 Collection"""
    connect_milvus()
    if not utility.has_collection(collection_name):
        raise KnowledgeProcessingError("Milvus 集合不存在")
    fields = {field.name: field for field in Collection(collection_name).schema.fields}
    primary = next((field for field in fields.values() if field.is_primary), None)
    if primary is None or primary.auto_id or "document_id" not in fields:
        raise LegacyCollectionSchemaError(
            "该知识库使用旧版 auto-ID Collection，无法保证幂等入库；"
            "请新建知识库并迁移文档后再上传"
        )


def stable_chunk_id(document_id: str, chunk_index: int) -> int:
    """利用 SHA256 为文档切片确定性地生成唯一正整数 INT64 主键"""
    digest = hashlib.sha256(f"{document_id}:{chunk_index}".encode()).digest()
    value = int.from_bytes(digest[:8], "big") & ((1 << 63) - 1)
    return value or 1


async def embed_texts(texts: list[str], vector_dimension: int) -> list[list[float]]:
    """批处理调用向量模型 API 为文本生成 embedding 向量

    Args:
        texts (list[str]): 文本块列表
        vector_dimension (int): 预期返回向量维度

    Returns:
        list[list[float]]: 生成的浮点向量列表
    """
    api_key = os.getenv("OPENAI_API_KEY", "").strip() or os.getenv("DASHSCOPE_API_KEY", "").strip()
    api_base = os.getenv("OPENAI_API_BASE", "").strip().rstrip("/")
    model = os.getenv("EMBEDDING_MODEL", "").strip()
    if not api_key or not api_base or not model:
        raise EmbeddingServiceError("后端 .env 缺少 Embedding API Key、URL 或模型配置")

    vectors: list[list[float]] = []
    timeout = httpx.Timeout(connect=20.0, read=180.0, write=60.0, pool=20.0)
    async with httpx.AsyncClient(timeout=timeout) as client:
        for start in range(0, len(texts), 10):
            batch = texts[start : start + 10]
            response: httpx.Response | None = None
            try:
                async for attempt in AsyncRetrying(
                    stop=stop_after_attempt(3),
                    wait=wait_exponential(multiplier=0.5, min=0.5, max=4),
                    retry=retry_if_exception_type(
                        (httpx.TransportError, TransientEmbeddingError)
                    ),
                    reraise=True,
                ):
                    with attempt:
                        response = await client.post(
                            f"{api_base}/embeddings",
                            headers={
                                "Authorization": f"Bearer {api_key}",
                                "Content-Type": "application/json",
                            },
                            json={
                                "model": model,
                                "input": batch,
                                "dimensions": vector_dimension,
                                "encoding_format": "float",
                            },
                        )
                        if response.status_code == 429 or response.status_code >= 500:
                            raise TransientEmbeddingError(
                                f"向量服务暂时不可用: {response.status_code}"
                            )
            except (httpx.TransportError, TransientEmbeddingError) as exc:
                raise EmbeddingServiceError(str(exc)) from exc
            assert response is not None
            if response.status_code >= 400:
                try:
                    payload = response.json()
                    detail = payload.get("error", {}).get("message") or str(payload)
                except ValueError:
                    detail = response.text
                raise EmbeddingServiceError(
                    f"向量模型返回 {response.status_code}: {detail[:500]}"
                )
            payload = response.json()
            data = sorted(payload.get("data", []), key=lambda item: item.get("index", 0))
            batch_vectors = [item.get("embedding") for item in data]
            if len(batch_vectors) != len(batch) or any(not isinstance(vector, list) for vector in batch_vectors):
                raise EmbeddingServiceError("向量模型返回的数据数量或格式不正确")
            for vector in batch_vectors:
                if len(vector) != vector_dimension:
                    raise EmbeddingServiceError(
                        f"向量模型实际返回 {len(vector)} 维，与知识库设定的 {vector_dimension} 维不一致"
                    )
            vectors.extend(batch_vectors)
    return vectors


def insert_chunks(
    collection_name: str,
    chunks: list[dict[str, Any]],
    vectors: list[list[float]],
    document_name: str,
    source_type: str,
    uploaded_by: int,
    document_id: str,
) -> list[int]:
    """将文档切片及向量数据 upsert 覆盖插入 Milvus Collection

    Returns:
        list[int]: 生成的稳定主键 ID 列表
    """
    connect_milvus()
    require_idempotent_collection(collection_name)
    if len(chunks) != len(vectors):
        raise KnowledgeProcessingError("切片数量与向量数量不一致")
    collection = Collection(collection_name)
    timestamp = datetime.now(timezone.utc).isoformat()
    primary_keys = [
        stable_chunk_id(document_id, index) for index in range(len(chunks))
    ]
    collection.upsert(
        [
            {
                "id": primary_keys[index],
                "document_id": document_id,
                "text": chunk["text"],
                "document_name": document_name,
                "source_type": source_type,
                "chunk_index": index,
                "token_count": chunk["token_count"],
                "uploaded_by": uploaded_by,
                "created_at": timestamp,
                "embedding": vector,
            }
            for index, (chunk, vector) in enumerate(zip(chunks, vectors))
        ]
    )
    collection.flush()
    return primary_keys


def delete_chunks(collection_name: str, primary_keys: list[int]) -> None:
    """按主键列表批量从 Milvus 中删除指定的切片向量"""
    if not primary_keys:
        return
    connect_milvus()
    collection = Collection(collection_name)
    ids = ", ".join(str(int(primary_key)) for primary_key in primary_keys)
    collection.delete(f"id in [{ids}]")
    collection.flush()


def _live_entity_count(collection: Collection) -> int:
    """获取集合中实时存活的实体记录数量"""
    collection.load()
    result = collection.query(expr="", output_fields=["count(*)"])
    return int(result[0]["count(*)"]) if result else 0


def collection_details(collection_name: str) -> dict[str, Any]:
    """获取 Milvus 集合的详细结构、字段及索引列表"""
    connect_milvus()
    if not utility.has_collection(collection_name):
        raise KnowledgeProcessingError("Milvus 集合不存在")
    collection = Collection(collection_name)
    fields = []
    for field in collection.schema.fields:
        params = dict(field.params)
        fields.append(
            {
                "name": field.name,
                "type": field.dtype.name,
                "is_primary": field.is_primary,
                "auto_id": field.auto_id,
                "dimension": params.get("dim"),
                "max_length": params.get("max_length"),
            }
        )
    return {
        "collection_name": collection_name,
        "description": collection.schema.description,
        "entity_count": _live_entity_count(collection),
        "fields": fields,
        "indexes": [index.to_dict() for index in collection.indexes],
    }


def list_chunks(
    collection_name: str,
    offset: int = 0,
    limit: int = 50,
    cursor: int | None = None,
) -> dict[str, Any]:
    """游标/分页方式列出 Milvus 集合中已存储的文档切片列表"""
    connect_milvus()
    if not utility.has_collection(collection_name):
        raise KnowledgeProcessingError("Milvus 集合不存在")
    collection = Collection(collection_name)
    collection.load()
    total = _live_entity_count(collection)
    output_fields = [
        "id",
        "text",
        "document_name",
        "source_type",
        "chunk_index",
        "token_count",
        "uploaded_by",
        "created_at",
    ]
    if any(field.name == "document_id" for field in collection.schema.fields):
        output_fields.append("document_id")
    rows = collection.query(
        expr=f"id > {int(cursor)}" if cursor is not None else "id >= 0",
        output_fields=output_fields,
        offset=0 if cursor is not None else max(0, offset),
        limit=max(1, min(limit, 200)),
    )
    rows.sort(key=lambda row: row.get("id", 0))
    next_cursor = int(rows[-1]["id"]) if len(rows) == limit else None
    return {
        "items": rows,
        "offset": offset,
        "limit": limit,
        "total": total,
        "next_cursor": next_cursor,
    }
